#!/usr/bin/env python3
"""類題をWordファイルとして出力するスクリプト"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_border(cell, **kwargs):
    """セルの罫線を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            from lxml import etree
            element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))

def add_tile_row_to_cell(cell, pattern, cell_size=18):
    """セルにタイルパターンを描画するためのテーブルを追加"""
    # テキストで表現（■□）
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(''.join(['■' if x else '□' for x in pattern]))
    run.font.size = Pt(14)
    run.font.name = 'MS Gothic'
    # フォールバック用にeast asiaフォントも設定
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), 'MS Gothic')

def add_tile_pattern(doc, patterns, label=None):
    """タイルパターン（4行）を文書に追加"""
    if label:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.name = 'MS Gothic'

    for i, pattern in enumerate(patterns):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        text = '  '.join(['■' if x else '□' for x in pattern])
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = 'MS Gothic'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), 'MS Gothic')

        # 矢印を追加（最後の行以外）
        if i < len(patterns) - 1:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            run2 = p2.add_run('      ↓')
            run2.font.size = Pt(10)


def create_document():
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ====================
    # 問1
    # ====================
    p = doc.add_paragraph()
    run = p.add_run('問 1')
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        '図のように、黒と白に色が変化するタイルが、ある規則にしたがって、'
        '上から下のように変化している。'
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        '図と同じ規則にしたがって変化しているタイルは次のうちどれか。'
        '正しいものを1つ選び、解答用紙にマークせよ。'
        'ただし、選択肢では矢印は省略している。'
    )
    run.font.size = Pt(11)

    # 例題パターン
    example = [
        [1, 0, 1, 1, 0, 1],
        [1, 1, 1, 0, 1, 1],
        [1, 0, 0, 1, 1, 0],
        [1, 1, 0, 1, 0, 1],
    ]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('【例】')
    run.font.size = Pt(11)
    run.bold = True

    add_tile_pattern(doc, example)

    doc.add_paragraph()  # 空行

    # 選択肢
    options_data = [
        # (1)
        [
            [0, 1, 1, 0, 1, 0],
            [0, 0, 1, 1, 0, 1],
            [0, 0, 0, 1, 1, 0],
            [0, 0, 0, 0, 1, 1],
        ],
        # (2)
        [
            [1, 0, 0, 1, 1, 0],
            [1, 1, 0, 1, 0, 1],
            [1, 0, 1, 0, 1, 1],
            [1, 1, 1, 1, 0, 0],
        ],
        # (3)
        [
            [0, 1, 0, 1, 1, 0],
            [0, 1, 1, 1, 0, 1],
            [0, 1, 0, 0, 1, 0],
            [0, 1, 1, 0, 1, 1],
        ],
        # (4) 正解
        [
            [1, 1, 0, 1, 0, 0],
            [1, 0, 1, 1, 1, 0],
            [1, 1, 1, 0, 0, 1],
            [1, 0, 0, 1, 0, 1],
        ],
        # (5)
        [
            [1, 0, 1, 0, 1, 1],
            [1, 1, 1, 1, 1, 0],
            [1, 0, 0, 0, 0, 1],
            [1, 1, 1, 0, 0, 1],  # wrong (correct would be ■■□□■■)
        ],
    ]

    # 選択肢を2列+3列で表示するためにテーブルを使用
    # まず (1)〜(5) のラベル付きで表示
    for idx, opt in enumerate(options_data):
        label = f'（{idx + 1}）'
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(label)
        run.font.size = Pt(11)

        for row in opt:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1)
            text = '  '.join(['■' if x else '□' for x in row])
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = 'MS Gothic'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from lxml import etree
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), 'MS Gothic')

    # ページ区切り
    doc.add_page_break()

    # ====================
    # 問2
    # ====================
    p = doc.add_paragraph()
    run = p.add_run('問 2')
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        '立方体ABCD-EFGHがある（ABCDが上面、EFGHが底面で、'
        'A-E、B-F、C-G、D-Hがそれぞれ対応する頂点）。'
        'この立方体の8つの頂点のうち4つを選び、それらを頂点とする正四面体を作ることを考える。'
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'このとき、正四面体の1辺の長さは立方体の1辺の長さの何倍か。'
        '正しいものを1つ選び、解答用紙にマークせよ。'
    )
    run.font.size = Pt(11)

    # 立方体の図（テキストで説明）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 立方体のASCII図
    cube_lines = [
        '          C ─────── D',
        '         /|        /|',
        '        / |       / |',
        '       B ─────── A  |',
        '       |  G ─ ─ ─|─ H',
        '       | /        | /',
        '       |/         |/',
        '       F ─────── E',
    ]

    for line in cube_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = 'Courier New'

    doc.add_paragraph()

    # 選択肢 (正解は(3) √2倍)
    choices = [
        '（1）  1 倍',
        '（2）  √3 / √2 倍',
        '（3）  √2 倍',
        '（4）  √3 倍',
        '（5）  2 倍',
    ]

    for choice in choices:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(choice)
        run.font.size = Pt(12)

    # ページ区切り
    doc.add_page_break()

    # ====================
    # 解説ページ
    # ====================
    p = doc.add_paragraph()
    run = p.add_run('解 説')
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 問1 解説
    p = doc.add_paragraph()
    run = p.add_run('問1（正解: (4)）')
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        '規則: 各行から次の行への変化は、隣接するセルのXOR演算による。\n'
        '具体的には、新しい行の各セルは以下のように決まる:\n'
        '・最左端のセル: 変化しない（前の行と同じ）\n'
        '・それ以外のセルi: 前の行のセル(i-1)とセルiのXOR\n'
        '  （同じ色→白、異なる色→黒）\n\n'
        '例題で確認:\n'
        '行1: ■□■■□■\n'
        '行2: ■■■□■■  （■→■, □XOR■=■, ■XOR■=□, ...）\n\n'
        '選択肢(4)のみがこの規則に完全に従っている。'
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    # 問2 解説
    p = doc.add_paragraph()
    run = p.add_run('問2（正解: (3) √2倍）')
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        '立方体の8頂点から4つを選んで正四面体を作るには、\n'
        '対角線上にある4頂点を選ぶ。\n\n'
        '例えば、A, C, F, H の4頂点を選ぶと:\n'
        '・ACは上面の対角線（面の対角線）\n'
        '・AFは側面の対角線（面の対角線）\n'
        '・AHは底面方向の対角線（面の対角線）\n\n'
        '立方体の1辺の長さをaとすると、\n'
        '面の対角線の長さ = √(a² + a²) = a√2\n\n'
        'A, C, F, H のどの2点間の距離も a√2 で等しいため、\n'
        'これらは正四面体をなす。\n\n'
        'したがって、正四面体の1辺の長さは立方体の1辺の √2 倍。\n\n'
        '（同様に B, D, E, G でも正四面体が作れる。'
        '正四面体は2組のみ。）'
    )
    run.font.size = Pt(10)

    # 保存
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '類題.docx')
    doc.save(output_path)
    print(f"保存完了: {output_path}")
    return output_path

if __name__ == '__main__':
    create_document()
